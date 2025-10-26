#!/usr/bin/env python3
"""
Analyze DOCX structure to understand content before first heading.
"""

import sys
from docx import Document
from docx.oxml.ns import qn


def is_heading(para, level=1):
    """Check if paragraph is a heading."""
    heading_style = f"Heading {level}"

    if para.style and para.style.name == heading_style:
        return True

    if hasattr(para, '_element'):
        pPr = para._element.pPr
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                outline_level = int(outlineLvl.get(qn('w:val')))
                if outline_level == level - 1:
                    return True

    return False


def main():
    if len(sys.argv) < 2:
        print("Usage: python analyze_docx.py <file.docx>")
        sys.exit(1)

    filename = sys.argv[1]
    doc = Document(filename)

    print(f"Analyzing: {filename}")
    print(f"Total paragraphs: {len(doc.paragraphs)}\n")

    # Find first Heading 1
    first_heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        if is_heading(para, 1):
            first_heading_idx = i
            print(f"First Heading 1 found at paragraph {i}: '{para.text}'")
            break

    if first_heading_idx is None:
        print("No Heading 1 found!")
        return

    # Show content before first heading
    if first_heading_idx > 0:
        print(f"\n⚠️  Found {first_heading_idx} paragraphs BEFORE first Heading 1:")
        print("=" * 70)
        for i in range(first_heading_idx):
            para = doc.paragraphs[i]
            style = para.style.name if para.style else "None"
            text = para.text[:80] if para.text else "(empty)"
            print(f"  {i}: [{style}] {text}")
        print("=" * 70)
        print(f"\nThese {first_heading_idx} paragraphs will NOT be captured by the split script!")
        print("This is a BUG that needs to be fixed.\n")
    else:
        print("\n✓ No content before first Heading 1")

    # Count headings
    heading_count = sum(1 for para in doc.paragraphs if is_heading(para, 1))
    print(f"\nTotal Heading 1 count: {heading_count}")


if __name__ == '__main__':
    main()
