#!/usr/bin/env python3
"""
Simplify DOCX styles by consolidating them into a canonical set.

This script analyzes all styles used in a document and maps them to
a minimal canonical set, one style per element type.
"""

import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from collections import defaultdict, Counter


class StyleSimplifier:
    """Simplifies document styles to a canonical minimal set."""

    # Canonical style mapping - maps style patterns to canonical styles
    STYLE_MAPPING = {
        # Headings
        'Heading 1': 'Heading 1',
        'Heading 2': 'Heading 2',
        'Heading 3': 'Heading 3',
        'Heading 4': 'Heading 3',  # Map H4+ to H3
        'Heading 5': 'Heading 3',
        'Heading 6': 'Heading 3',
        'Heading 7': 'Heading 3',
        'Heading 8': 'Heading 3',
        'Heading 9': 'Heading 3',

        # Title elements
        'Title': 'Title',
        'Subtitle': 'Subtitle',

        # Body text
        'Normal': 'Normal',
        'Body Text': 'Normal',
        'Body Text 2': 'Normal',
        'Body Text 3': 'Normal',
        'First Paragraph': 'Normal',

        # Lists
        'List': 'Normal',
        'List Paragraph': 'Normal',
        'List Bullet': 'Normal',
        'List Number': 'Normal',
        'List Continue': 'Normal',

        # Special elements
        'Quote': 'Quote',
        'Block Quote': 'Quote',
        'Intense Quote': 'Quote',
        'Caption': 'Caption',
        'Footer': 'Caption',
        'Header': 'Caption',
    }

    # Canonical style set we want to keep
    CANONICAL_STYLES = {
        'Title',      # Document title
        'Subtitle',   # Document subtitle
        'Heading 1',  # Major sections (Acts, etc.)
        'Heading 2',  # Chapters
        'Heading 3',  # Subsections within chapters
        'Normal',     # Body text
        'Quote',      # Block quotes
        'Caption',    # Captions, headers, footers
    }

    def __init__(self, input_file, output_file=None):
        """Initialize the simplifier."""
        self.input_file = input_file
        self.output_file = output_file or self._generate_output_filename()

    def _generate_output_filename(self):
        """Generate output filename from input."""
        path = Path(self.input_file)
        return str(path.parent / f"{path.stem}_simplified{path.suffix}")

    def analyze(self):
        """Analyze the document styles and show what would be simplified."""
        print(f"Analyzing: {self.input_file}")
        print("=" * 80)

        doc = Document(self.input_file)

        # Count style usage
        style_usage = Counter()
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else 'None'
            style_usage[style_name] += 1

        # Group by canonical mapping
        canonical_groups = defaultdict(list)
        unmapped_styles = []

        for style_name, count in style_usage.most_common():
            canonical = self._get_canonical_style(style_name)
            if canonical:
                canonical_groups[canonical].append((style_name, count))
            else:
                unmapped_styles.append((style_name, count))

        # Display analysis
        print(f"\nTotal paragraphs: {sum(style_usage.values())}")
        print(f"Unique styles: {len(style_usage)}")
        print(f"Would consolidate to: {len(canonical_groups)} canonical styles")

        print("\n" + "=" * 80)
        print("CANONICAL STYLE MAPPING:")
        print("=" * 80)

        for canonical in sorted(canonical_groups.keys()):
            styles = canonical_groups[canonical]
            total_paras = sum(count for _, count in styles)
            print(f"\n{canonical} ({total_paras} paragraphs):")
            for style_name, count in sorted(styles, key=lambda x: -x[1]):
                if style_name == canonical:
                    print(f"  • {style_name}: {count} paragraphs (canonical)")
                else:
                    print(f"  • {style_name}: {count} paragraphs → will map to {canonical}")

        if unmapped_styles:
            print(f"\n" + "=" * 80)
            print(f"UNMAPPED STYLES (will use default 'Normal'):")
            print("=" * 80)
            for style_name, count in unmapped_styles:
                print(f"  • {style_name}: {count} paragraphs")

        return canonical_groups, unmapped_styles

    def _get_canonical_style(self, style_name):
        """Get the canonical style for a given style name."""
        if not style_name or style_name == 'None':
            return 'Normal'

        # Direct mapping
        if style_name in self.STYLE_MAPPING:
            return self.STYLE_MAPPING[style_name]

        # Pattern matching
        style_lower = style_name.lower()

        if 'heading' in style_lower:
            # Extract heading number if present
            import re
            match = re.search(r'(\d+)', style_name)
            if match:
                level = int(match.group(1))
                if level <= 3:
                    return f'Heading {level}'
                else:
                    return 'Heading 3'
            return 'Heading 1'

        if any(word in style_lower for word in ['title', 'heading']):
            return 'Title'

        if any(word in style_lower for word in ['quote', 'quotation']):
            return 'Quote'

        if any(word in style_lower for word in ['caption', 'footer', 'header']):
            return 'Caption'

        # Default to Normal
        return 'Normal'

    def simplify(self, dry_run=False):
        """Simplify the document styles."""
        print(f"\n{'DRY RUN: ' if dry_run else ''}Simplifying styles in: {self.input_file}")
        print("=" * 80)

        doc = Document(self.input_file)

        # Track changes
        changes = Counter()

        # Apply canonical styles to paragraphs
        for para in doc.paragraphs:
            original_style = para.style.name if para.style else 'None'
            canonical_style = self._get_canonical_style(original_style)

            if original_style != canonical_style:
                changes[f"{original_style} → {canonical_style}"] += 1
                if not dry_run:
                    try:
                        para.style = canonical_style
                    except KeyError:
                        # Style doesn't exist, use Normal
                        para.style = 'Normal'
                        changes[f"{original_style} → Normal (fallback)"] += 1

        # Display changes
        print(f"\nStyle changes:")
        for change, count in sorted(changes.items(), key=lambda x: -x[1]):
            print(f"  • {change}: {count} paragraphs")

        if not dry_run:
            # Save simplified document
            doc.save(self.output_file)
            print(f"\n✓ Saved simplified document to: {self.output_file}")
        else:
            print(f"\n(Dry run - no changes made)")

        return len(changes)


def main():
    parser = argparse.ArgumentParser(
        description="Simplify DOCX styles to a canonical minimal set",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Analyze styles (no changes)
  %(prog)s mybook.docx --analyze

  # Dry run (show what would change)
  %(prog)s mybook.docx --dry-run

  # Simplify and save to new file
  %(prog)s mybook.docx

  # Simplify and specify output file
  %(prog)s mybook.docx --output simplified.docx

Canonical styles:
  • Title       - Document title
  • Subtitle    - Document subtitle
  • Heading 1   - Major sections (Acts, Parts)
  • Heading 2   - Chapters
  • Heading 3   - Subsections
  • Normal      - Body text, lists, etc.
  • Quote       - Block quotes
  • Caption     - Captions, headers, footers
        """
    )

    parser.add_argument('input_file', help='Input DOCX file')
    parser.add_argument('-o', '--output', help='Output file (default: INPUT_simplified.docx)')
    parser.add_argument('-a', '--analyze', action='store_true',
                        help='Analyze only, do not modify')
    parser.add_argument('-d', '--dry-run', action='store_true',
                        help='Dry run - show changes but do not save')

    args = parser.parse_args()

    # Check input file exists
    if not Path(args.input_file).exists():
        print(f"ERROR: File not found: {args.input_file}")
        sys.exit(1)

    # Create simplifier
    output_file = args.output if args.output else None
    simplifier = StyleSimplifier(args.input_file, output_file)

    if args.analyze:
        # Just analyze and show mapping
        simplifier.analyze()
    else:
        # Show analysis first
        simplifier.analyze()

        # Then simplify
        print("\n" + "=" * 80)
        simplifier.simplify(dry_run=args.dry_run)


if __name__ == '__main__':
    main()
