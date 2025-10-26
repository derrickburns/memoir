#!/usr/bin/env python3
"""
Split a DOCX book into chapters with image extraction.

This script:
- Reads a DOCX file
- Splits it into chapters based on heading styles
- Extracts images into chapter-specific directories
- Creates one DOCX file per chapter
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import argparse


class BookSplitter:
    def __init__(self, input_file, output_dir="output", heading_level=1):
        """
        Initialize the book splitter.

        Args:
            input_file: Path to the input DOCX file
            output_dir: Directory where output files will be saved
            heading_level: Heading level to use for chapter detection (1-9)
        """
        self.input_file = input_file
        self.output_dir = Path(output_dir)
        self.heading_style = f"Heading {heading_level}"
        self.heading_level = heading_level

        # Create output directory
        self.output_dir.mkdir(exist_ok=True)

    def extract_image(self, run, chapter_num, image_num, images_dir):
        """Extract an image from a run element."""
        images_dir.mkdir(parents=True, exist_ok=True)

        # Get the image data
        drawing_elements = run._element.findall(f'.//{qn("a:blip")}',
                                                 namespaces=run._element.nsmap)

        for drawing in drawing_elements:
            embed = drawing.get(qn('r:embed'))
            if embed:
                try:
                    image_part = run.part.related_parts[embed]
                    image_bytes = image_part.blob

                    # Determine file extension
                    content_type = image_part.content_type
                    ext = content_type.split('/')[-1]
                    if ext == 'jpeg':
                        ext = 'jpg'

                    # Save image
                    image_filename = f"image_{image_num:03d}.{ext}"
                    image_path = images_dir / image_filename

                    with open(image_path, 'wb') as f:
                        f.write(image_bytes)

                    print(f"  Extracted image: {image_path}")
                    return image_filename
                except KeyError:
                    # Image relationship doesn't exist (happens in reconstructed docs)
                    # Images are still embedded, just can't extract them separately
                    return None

        return None

    def has_image(self, run):
        """Check if a run contains an image."""
        drawing_elements = run._element.findall(f'.//{qn("a:blip")}',
                                                 namespaces=run._element.nsmap)
        return len(drawing_elements) > 0

    def extract_chapter_number(self, title):
        """
        Extract chapter number from title.

        Returns the chapter number if found, None otherwise.
        Handles formats like: "Chapter 2", "Chapter 10", "CHAPTER 5", etc.
        """
        patterns = [
            r'[Cc]hapter\s+(\d+)',
            r'[Cc]h\.\s*(\d+)',
            r'^(\d+)\.',
            r'^(\d+)\s',
        ]

        for pattern in patterns:
            match = re.search(pattern, title)
            if match:
                return int(match.group(1))

        return None

    def is_chapter_heading(self, paragraph):
        """Check if a paragraph is a chapter heading."""
        # Check by style name
        if paragraph.style.name == self.heading_style:
            return True

        # Also check by outline level for robustness
        if hasattr(paragraph, '_element'):
            pPr = paragraph._element.pPr
            if pPr is not None:
                outlineLvl = pPr.find(qn('w:outlineLvl'))
                if outlineLvl is not None:
                    level = int(outlineLvl.get(qn('w:val')))
                    if level == self.heading_level - 1:  # Heading 1 = outline level 0
                        return True

        return False

    def split(self):
        """Split the document into chapters."""
        print(f"Reading document: {self.input_file}")
        doc = Document(self.input_file)

        chapters = []
        current_chapter = None
        sequential_num = 0
        front_matter = []  # Collect content before first heading

        # Group paragraphs into chapters
        for para in doc.paragraphs:
            if self.is_chapter_heading(para):
                # Skip empty headings (they're formatting artifacts)
                if not para.text.strip():
                    continue

                # Start a new chapter
                if current_chapter is not None:
                    chapters.append(current_chapter)

                sequential_num += 1
                chapter_title = para.text.strip()

                # Extract actual chapter number from title if present
                chapter_number = self.extract_chapter_number(chapter_title)

                current_chapter = {
                    'number': chapter_number,  # None if not found, int if found
                    'title': chapter_title,
                    'elements': [para],
                    'sequential': sequential_num  # For reference
                }

                if chapter_number:
                    print(f"Found chapter {chapter_number}: {chapter_title}")
                else:
                    print(f"Found section: {chapter_title}")
            elif current_chapter is not None:
                current_chapter['elements'].append(para)
            else:
                # Content before first heading
                front_matter.append(para)

        # Don't forget the last chapter
        if current_chapter is not None:
            chapters.append(current_chapter)

        if not chapters:
            print("ERROR: No chapters found! Make sure your document has headings.")
            print(f"Looking for style: '{self.heading_style}'")
            print("\nAvailable styles in document:")
            for style in doc.styles:
                if 'Heading' in style.name:
                    print(f"  - {style.name}")
            return

        print(f"\nFound {len(chapters)} chapters total")

        # Save front matter if it exists
        if front_matter:
            print(f"\nFound {len(front_matter)} paragraphs of front matter (before first heading)")
            front_matter_chapter = {
                'number': None,
                'title': '00_Front_Matter',
                'elements': front_matter,
                'sequential': 0
            }
            self.save_chapter(doc, front_matter_chapter)

        # Process each chapter
        for chapter in chapters:
            self.save_chapter(doc, chapter)

        print(f"\nDone! Output saved to: {self.output_dir}")

    def save_chapter(self, original_doc, chapter):
        """Save a chapter as a separate DOCX file with images extracted."""
        chapter_num = chapter['number']  # Can be None or int
        chapter_title = chapter['title']
        sequential_num = chapter['sequential']

        # Create sanitized filename
        safe_title = "".join(c for c in chapter_title if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_title = safe_title.replace(' ', '_')

        # Use actual chapter number if available, otherwise just use the title
        if chapter_num is not None:
            chapter_filename = f"chapter_{chapter_num:02d}_{safe_title}.docx"
            images_dir_name = f"chapter_{chapter_num:02d}_images"
            display_name = f"Chapter {chapter_num}: {chapter_title}"
        else:
            chapter_filename = f"{safe_title}.docx"
            images_dir_name = f"{safe_title}_images"
            display_name = chapter_title

        chapter_path = self.output_dir / chapter_filename
        images_dir = self.output_dir / images_dir_name

        print(f"\nProcessing {display_name}")

        # Create new document as a clone of the original (to preserve styles)
        import copy
        from docx.oxml.document import CT_Document

        # Create a new document with the same styles as the original
        chapter_doc = Document()

        # Properly copy all styles by copying the entire styles part
        chapter_doc.styles._element.clear()
        for style_element in original_doc.styles._element:
            chapter_doc.styles._element.append(copy.deepcopy(style_element))

        image_counter = 1

        # Add content to chapter document by deep copying entire paragraphs
        for element in chapter['elements']:
            # Deep copy the entire paragraph element to preserve everything:
            # - style (the pStyle attribute will reference styles we copied above)
            # - formatting
            # - runs
            # - images
            # - all properties
            new_para_element = copy.deepcopy(element._element)
            chapter_doc._element.body.append(new_para_element)

            # Count images for extraction
            for run in element.runs:
                if self.has_image(run):
                    self.extract_image(run, chapter_num, image_counter, images_dir)
                    image_counter += 1

        # Save chapter document
        chapter_doc.save(chapter_path)
        print(f"  Saved: {chapter_path}")

        if image_counter > 1:
            print(f"  Images: {image_counter - 1} images saved to {images_dir}")


def main():
    parser = argparse.ArgumentParser(
        description="Split a DOCX book into chapters with image extraction",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s mybook.docx
  %(prog)s mybook.docx --output chapters --heading 2
  %(prog)s mybook.docx -o output -l 1
        """
    )

    parser.add_argument('input_file', help='Input DOCX file to split')
    parser.add_argument('-o', '--output', default='output',
                        help='Output directory (default: output)')
    parser.add_argument('-l', '--heading-level', type=int, default=1,
                        choices=range(1, 10),
                        help='Heading level for chapter detection (1-9, default: 1)')

    args = parser.parse_args()

    # Check if input file exists
    if not os.path.exists(args.input_file):
        print(f"ERROR: Input file not found: {args.input_file}")
        sys.exit(1)

    # Check if input file is DOCX
    if not args.input_file.lower().endswith('.docx'):
        print(f"ERROR: Input file must be a .docx file")
        sys.exit(1)

    # Create splitter and run
    splitter = BookSplitter(args.input_file, args.output, args.heading_level)
    splitter.split()


if __name__ == '__main__':
    main()
