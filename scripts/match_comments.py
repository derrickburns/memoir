#!/usr/bin/env python3
"""
Match extracted comments to split chapter files.

Reads comments.json and finds the corresponding text in the chapter DOCX files,
creating a mapping of comments to chapters.
"""

import json
import sys
from pathlib import Path
from docx import Document
from difflib import SequenceMatcher
import re


class CommentMatcher:
    def __init__(self, comments_file='comments.json', chapters_dir='output'):
        """Initialize the matcher."""
        self.comments_file = comments_file
        self.chapters_dir = Path(chapters_dir)
        self.comments_data = None
        self.chapter_texts = {}

    def load_comments(self):
        """Load comments from JSON file."""
        print(f"Loading comments from {self.comments_file}...")

        if not Path(self.comments_file).exists():
            print(f"ERROR: {self.comments_file} not found!")
            print("\nRun extract_comments.py first to extract comments from Google Doc")
            sys.exit(1)

        with open(self.comments_file, 'r') as f:
            self.comments_data = json.load(f)

        print(f"✓ Loaded {self.comments_data['total_comments']} comments")
        return self.comments_data

    def load_chapter_texts(self):
        """Load all chapter files and extract their text."""
        print(f"\nLoading chapter files from {self.chapters_dir}...")

        chapter_files = sorted(self.chapters_dir.glob('*.docx'))

        for chapter_file in chapter_files:
            # Skip temporary Word files
            if chapter_file.name.startswith('~$'):
                continue

            try:
                doc = Document(chapter_file)
                text = '\n'.join([para.text for para in doc.paragraphs])
                self.chapter_texts[chapter_file.name] = {
                    'path': str(chapter_file),
                    'text': text,
                    'paragraphs': [para.text for para in doc.paragraphs]
                }
            except Exception as e:
                print(f"Warning: Could not read {chapter_file.name}: {e}")

        print(f"✓ Loaded {len(self.chapter_texts)} chapter files")

    def normalize_text(self, text):
        """Normalize text for matching (remove extra whitespace, etc.)."""
        # Replace multiple spaces/newlines with single space
        text = re.sub(r'\s+', ' ', text)
        # Remove leading/trailing whitespace
        text = text.strip()
        return text

    def find_text_in_chapters(self, search_text, threshold=0.6):
        """
        Find which chapter(s) contain the given text.

        Returns a list of matches with similarity scores.
        """
        if not search_text or not search_text.strip():
            return []

        search_text_norm = self.normalize_text(search_text)
        matches = []

        for chapter_name, chapter_data in self.chapter_texts.items():
            chapter_text_norm = self.normalize_text(chapter_data['text'])

            # Try exact match first
            if search_text_norm in chapter_text_norm:
                # Find the position
                pos = chapter_text_norm.index(search_text_norm)
                # Get context (50 chars before and after)
                context_start = max(0, pos - 50)
                context_end = min(len(chapter_text_norm), pos + len(search_text_norm) + 50)
                context = chapter_text_norm[context_start:context_end]

                matches.append({
                    'chapter': chapter_name,
                    'similarity': 1.0,
                    'match_type': 'exact',
                    'context': context,
                    'position': pos
                })
            else:
                # Try fuzzy matching
                similarity = SequenceMatcher(None, search_text_norm, chapter_text_norm).ratio()

                if similarity > threshold:
                    matches.append({
                        'chapter': chapter_name,
                        'similarity': similarity,
                        'match_type': 'fuzzy',
                        'context': search_text_norm[:100] + '...',
                        'position': -1
                    })

        # Sort by similarity (highest first)
        matches.sort(key=lambda x: x['similarity'], reverse=True)
        return matches

    def match_all_comments(self):
        """Match all comments to chapters."""
        print("\nMatching comments to chapters...")

        results = {
            'document_title': self.comments_data['document_title'],
            'total_comments': self.comments_data['total_comments'],
            'matched_comments': 0,
            'unmatched_comments': 0,
            'comment_matches': []
        }

        for i, comment in enumerate(self.comments_data['comments'], 1):
            print(f"Processing comment {i}/{results['total_comments']}...", end='\r')

            quoted_text = comment.get('quoted_text', '')
            matches = self.find_text_in_chapters(quoted_text)

            comment_match = {
                'comment_id': comment['id'],
                'author': comment['author'],
                'content': comment['content'],
                'quoted_text': quoted_text,
                'created': comment['created'],
                'resolved': comment['resolved'],
                'matches': matches
            }

            if matches:
                results['matched_comments'] += 1
            else:
                results['unmatched_comments'] += 1

            results['comment_matches'].append(comment_match)

        print()  # New line after progress
        print(f"✓ Matched {results['matched_comments']} comments")
        print(f"  {results['unmatched_comments']} comments could not be matched")

        return results

    def export_matches(self, results, output_file='comment_matches.json'):
        """Export comment matches to JSON and markdown."""
        print(f"\nExporting matches to {output_file}...")

        # Save JSON
        with open(output_file, 'w') as f:
            json.dump(results, f, indent=2)

        print(f"✓ Saved to {output_file}")

        # Create markdown report
        md_file = output_file.replace('.json', '.md')
        with open(md_file, 'w') as f:
            f.write(f"# Comment Matches: {results['document_title']}\n\n")
            f.write(f"**Total comments:** {results['total_comments']}\n\n")
            f.write(f"**Matched:** {results['matched_comments']}\n\n")
            f.write(f"**Unmatched:** {results['unmatched_comments']}\n\n")
            f.write("---\n\n")

            # Group by chapter
            by_chapter = {}
            unmatched = []

            for comment_match in results['comment_matches']:
                if comment_match['matches']:
                    # Use the best match (first one, since sorted by similarity)
                    best_match = comment_match['matches'][0]
                    chapter = best_match['chapter']

                    if chapter not in by_chapter:
                        by_chapter[chapter] = []

                    by_chapter[chapter].append(comment_match)
                else:
                    unmatched.append(comment_match)

            # Write by chapter
            f.write("## Comments by Chapter\n\n")

            for chapter in sorted(by_chapter.keys()):
                f.write(f"### {chapter}\n\n")

                for comment in by_chapter[chapter]:
                    best_match = comment['matches'][0]
                    status = "✓ Resolved" if comment['resolved'] else "○ Open"

                    f.write(f"**{status}** - {comment['author']} ({comment['created']})\n\n")

                    if comment['quoted_text']:
                        f.write(f"> {comment['quoted_text'][:200]}")
                        if len(comment['quoted_text']) > 200:
                            f.write("...")
                        f.write("\n\n")

                    f.write(f"**Comment:** {comment['content']}\n\n")
                    f.write(f"*Match: {best_match['match_type']}, similarity: {best_match['similarity']:.2%}*\n\n")
                    f.write("---\n\n")

            # Write unmatched
            if unmatched:
                f.write("## Unmatched Comments\n\n")
                f.write("These comments could not be matched to any chapter:\n\n")

                for comment in unmatched:
                    status = "✓ Resolved" if comment['resolved'] else "○ Open"
                    f.write(f"**{status}** - {comment['author']} ({comment['created']})\n\n")
                    f.write(f"**Comment:** {comment['content']}\n\n")

                    if comment['quoted_text']:
                        f.write(f"**Quoted text:** {comment['quoted_text'][:100]}...\n\n")

                    f.write("---\n\n")

        print(f"✓ Created readable report: {md_file}")

        # Print summary
        print("\n" + "=" * 80)
        print("SUMMARY")
        print("=" * 80)
        print(f"\nComments by chapter:")
        for chapter in sorted(by_chapter.keys()):
            print(f"  {chapter}: {len(by_chapter[chapter])} comment(s)")

        if unmatched:
            print(f"\n  Unmatched: {len(unmatched)} comment(s)")

        print()


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description='Match extracted comments to split chapter files',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Match comments to chapters in output/ directory
  %(prog)s

  # Use custom comments file and chapters directory
  %(prog)s --comments my_comments.json --chapters-dir my_chapters/

  # Adjust matching threshold (0.0-1.0, default 0.6)
  %(prog)s --threshold 0.8

This script reads the comments.json file created by extract_comments.py
and tries to find the quoted text in the split chapter files.
        """
    )

    parser.add_argument('-c', '--comments', default='data/comments.json',
                        help='Comments JSON file (default: data/comments.json)')
    parser.add_argument('-d', '--chapters-dir', default='output',
                        help='Directory with chapter DOCX files (default: output)')
    parser.add_argument('-o', '--output', default='data/comment_matches.json',
                        help='Output file (default: data/comment_matches.json)')
    parser.add_argument('-t', '--threshold', type=float, default=0.6,
                        help='Fuzzy match threshold 0.0-1.0 (default: 0.6)')

    args = parser.parse_args()

    # Create matcher
    matcher = CommentMatcher(args.comments, args.chapters_dir)

    # Load data
    matcher.load_comments()
    matcher.load_chapter_texts()

    # Match comments
    results = matcher.match_all_comments()

    # Export
    matcher.export_matches(results, args.output)

    print("✓ Done!")


if __name__ == '__main__':
    main()
