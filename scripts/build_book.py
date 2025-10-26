#!/usr/bin/env python3
"""
Build a book from a manifest file.

This script reads book_structure.yaml and assembles the final document
from individual chapter files according to the manifest.
"""

import sys
import yaml
from pathlib import Path
from docx import Document
import argparse
import copy


class BookBuilder:
    def __init__(self, manifest_file="book_structure.yaml", output_file="book.docx"):
        """Initialize the builder."""
        self.manifest_file = manifest_file
        self.output_file = output_file
        self.manifest = None

    def load_manifest(self):
        """Load and parse the manifest file."""
        print(f"Loading manifest: {self.manifest_file}")

        with open(self.manifest_file, 'r') as f:
            self.manifest = yaml.safe_load(f)

        return self.manifest

    def show_structure(self):
        """Display the book structure from manifest."""
        print("\n" + "=" * 80)
        print("BOOK STRUCTURE")
        print("=" * 80)

        # Book info
        book_info = self.manifest.get('book', {})
        print(f"\nTitle: {book_info.get('title', 'N/A')}")
        print(f"Author: {book_info.get('author', 'N/A')}")

        # Front matter
        front_matter = self.manifest.get('front_matter', [])
        if front_matter:
            print(f"\n{'Front Matter':=^80}")
            for item in front_matter:
                status = "✓" if item.get('include', True) else "✗"
                print(f"  {status} {item.get('title', 'N/A')}")

        # Timeline
        timeline = self.manifest.get('timeline', [])
        if timeline:
            print(f"\n{'Timeline Sections':=^80}")
            included = [item for item in timeline if item.get('include', True)]
            for i, item in enumerate(timeline, 1):
                status = "✓" if item.get('include', True) else "✗"
                print(f"  {status} {i}. {item.get('title', 'N/A')}")

            print(f"  → {len(included)} of {len(timeline)} sections included")

        # Chapters
        chapters = self.manifest.get('chapters', [])
        if chapters:
            print(f"\n{'Chapters':=^80}")
            included_chapters = [ch for ch in chapters if ch.get('include', True)]

            # Show chapter listing
            for item in chapters:
                if item.get('include', True):
                    num = item.get('number', '?')
                    title = item.get('title', 'N/A')
                    print(f"  ✓ Chapter {num}: {title}")
                else:
                    num = item.get('number', '?')
                    title = item.get('title', 'N/A')
                    print(f"  ✗ Chapter {num}: {title} (EXCLUDED)")

            print(f"\n  → {len(included_chapters)} of {len(chapters)} chapters included")

        print("\n" + "=" * 80)

    def build(self, show_progress=True):
        """Build the book according to the manifest."""
        print(f"\nBuilding book: {self.output_file}")
        print("=" * 80)

        merged_doc = None
        section_count = 0

        # Add front matter
        front_matter = self.manifest.get('front_matter', [])
        for item in front_matter:
            if not item.get('include', True):
                continue

            file_path = item['file']
            if show_progress:
                print(f"Adding front matter: {item.get('title', Path(file_path).name)}")

            section_doc = Document(file_path)

            if merged_doc is None:
                merged_doc = section_doc
            else:
                self._append_document(merged_doc, section_doc)

            section_count += 1

        # Add timeline sections
        timeline = self.manifest.get('timeline', [])
        for item in timeline:
            if not item.get('include', True):
                continue

            file_path = item['file']
            if show_progress:
                print(f"Adding timeline: {item.get('title', Path(file_path).name)}")

            section_doc = Document(file_path)

            if merged_doc is None:
                merged_doc = section_doc
            else:
                self._append_document(merged_doc, section_doc)

            section_count += 1

        # Add chapters
        chapters = self.manifest.get('chapters', [])
        chapter_num = 1

        for item in chapters:
            if not item.get('include', True):
                if show_progress:
                    print(f"Skipping chapter: {item.get('title', 'N/A')} (excluded in manifest)")
                continue

            file_path = item['file']
            title = item.get('title', 'N/A')

            if show_progress:
                print(f"Adding Chapter {chapter_num}: {title}")

            section_doc = Document(file_path)

            if merged_doc is None:
                merged_doc = section_doc
            else:
                self._append_document(merged_doc, section_doc)

            section_count += 1
            chapter_num += 1

        # Save the merged document
        print(f"\nSaving to: {self.output_file}")
        merged_doc.save(self.output_file)

        print(f"\n✓ Book built successfully!")
        print(f"  Total sections: {section_count}")
        print(f"  Output: {self.output_file}")

    def _append_document(self, base_doc, append_doc):
        """Append content from one document to another."""
        # Merge styles
        existing_style_ids = {s.style_id for s in base_doc.styles}
        for style_element in append_doc.styles._element:
            style_id = style_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
            if style_id and style_id not in existing_style_ids:
                base_doc.styles._element.append(copy.deepcopy(style_element))
                existing_style_ids.add(style_id)

        # Copy all paragraphs
        for para in append_doc.paragraphs:
            new_para_element = copy.deepcopy(para._element)
            base_doc._element.body.append(new_para_element)


def main():
    parser = argparse.ArgumentParser(
        description="Build a book from a manifest file",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Show the book structure
  %(prog)s --show-structure

  # Build the book
  %(prog)s

  # Build with custom manifest and output
  %(prog)s --manifest my_book.yaml --output final.docx

Manifest Format:
  The manifest is a YAML file with sections:
  - book: metadata (title, author)
  - front_matter: list of front matter files
  - timeline: list of timeline sections
  - chapters: list of chapter files

  Each item can have:
  - file: path to the DOCX file
  - title: descriptive title
  - include: true/false (whether to include)
        """
    )

    parser.add_argument('-m', '--manifest', default='book_structure.yaml',
                        help='Manifest file (default: book_structure.yaml)')
    parser.add_argument('-o', '--output', default='book.docx',
                        help='Output file (default: book.docx)')
    parser.add_argument('-s', '--show-structure', action='store_true',
                        help='Show book structure and exit (no build)')
    parser.add_argument('-q', '--quiet', action='store_true',
                        help='Quiet mode - minimal output')

    args = parser.parse_args()

    # Check manifest exists
    if not Path(args.manifest).exists():
        print(f"ERROR: Manifest file not found: {args.manifest}")
        sys.exit(1)

    # Create builder
    builder = BookBuilder(args.manifest, args.output)

    # Load manifest
    try:
        builder.load_manifest()
    except Exception as e:
        print(f"ERROR loading manifest: {e}")
        sys.exit(1)

    # Show structure
    if not args.quiet:
        builder.show_structure()

    # Build if not just showing structure
    if not args.show_structure:
        print()
        builder.build(show_progress=not args.quiet)


if __name__ == '__main__':
    main()
