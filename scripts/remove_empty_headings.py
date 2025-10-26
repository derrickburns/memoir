#!/usr/bin/env python3
"""
Remove empty heading paragraphs from a DOCX file.

Empty headings are formatting artifacts that should be removed.
"""

import sys
import argparse
from pathlib import Path
from docx import Document


class EmptyHeadingRemover:
    def __init__(self, input_file, output_file=None):
        """Initialize the remover."""
        self.input_file = input_file
        self.output_file = output_file or self._generate_output_filename()

    def _generate_output_filename(self):
        """Generate output filename from input."""
        path = Path(self.input_file)
        return str(path.parent / f"{path.stem}_cleaned{path.suffix}")

    def analyze(self):
        """Find empty headings in the document."""
        print(f"Analyzing: {self.input_file}")
        print("=" * 80)

        doc = Document(self.input_file)

        empty_headings = []

        for i, para in enumerate(doc.paragraphs):
            if para.style and 'Heading' in para.style.name:
                if not para.text.strip():
                    empty_headings.append((i, para.style.name))

        if empty_headings:
            print(f"\nFound {len(empty_headings)} empty heading(s):")
            for para_num, style in empty_headings:
                print(f"  Paragraph {para_num}: [{style}] (empty)")
        else:
            print("\n✓ No empty headings found")

        return empty_headings

    def remove(self, dry_run=False):
        """Remove empty headings from the document."""
        print(f"\n{'DRY RUN: ' if dry_run else ''}Removing empty headings")
        print("=" * 80)

        doc = Document(self.input_file)

        # Find empty headings (need indices for removal)
        to_remove = []

        for i, para in enumerate(doc.paragraphs):
            if para.style and 'Heading' in para.style.name:
                if not para.text.strip():
                    to_remove.append(i)

        if not to_remove:
            print("No empty headings to remove")
            return 0

        print(f"Will remove {len(to_remove)} empty heading(s)")

        if not dry_run:
            # Remove paragraphs in reverse order to maintain indices
            for para_idx in reversed(to_remove):
                para = doc.paragraphs[para_idx]
                p = para._element
                p.getparent().remove(p)

            # Save the cleaned document
            doc.save(self.output_file)
            print(f"\n✓ Saved cleaned document to: {self.output_file}")
        else:
            print(f"\n(Dry run - no changes made)")

        return len(to_remove)


def main():
    parser = argparse.ArgumentParser(
        description="Remove empty heading paragraphs from DOCX files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Analyze only (no changes)
  %(prog)s mybook.docx --analyze

  # Dry run (show what would be removed)
  %(prog)s mybook.docx --dry-run

  # Remove empty headings
  %(prog)s mybook.docx

  # Specify output file
  %(prog)s mybook.docx --output clean.docx

This removes empty heading paragraphs that are formatting artifacts.
        """
    )

    parser.add_argument('input_file', help='Input DOCX file')
    parser.add_argument('-o', '--output', help='Output file (default: INPUT_cleaned.docx)')
    parser.add_argument('-a', '--analyze', action='store_true',
                        help='Analyze only, do not remove')
    parser.add_argument('-d', '--dry-run', action='store_true',
                        help='Dry run - show what would be removed')

    args = parser.parse_args()

    # Check input file exists
    if not Path(args.input_file).exists():
        print(f"ERROR: File not found: {args.input_file}")
        sys.exit(1)

    # Create remover
    output_file = args.output if args.output else None
    remover = EmptyHeadingRemover(args.input_file, output_file)

    # Analyze
    empty_headings = remover.analyze()

    if not args.analyze and empty_headings:
        print()
        remover.remove(dry_run=args.dry_run)


if __name__ == '__main__':
    main()
