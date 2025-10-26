#!/usr/bin/env python3
"""
Find all chapters in the document.
"""

import sys
import re
from docx import Document
from docx.oxml.ns import qn


def is_heading(para, level=1):
    """Check if paragraph is a heading."""
    heading_style = f"Heading {level}"

    if para.style and para.style.name == heading_style:
        return True

    if hasattr(para, '_element'):
        pPr = para._element.pPr
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                outline_level = int(outlineLvl.get(qn('w:val')))
                if outline_level == level - 1:
                    return True

    return False


def extract_chapter_number(text):
    """Extract chapter number from text."""
    patterns = [
        r'[Cc]hapter\s+(\d+)',
        r'[Cc]h\.\s*(\d+)',
        r'^(\d+)\.',
        r'^(\d+)\s',
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return int(match.group(1))

    return None


def main():
    if len(sys.argv) < 2:
        print("Usage: python find_chapters.py <file.docx>")
        sys.exit(1)

    filename = sys.argv[1]
    doc = Document(filename)

    print(f"Analyzing: {filename}")
    print(f"Total paragraphs: {len(doc.paragraphs)}\n")
    print("=" * 80)
    print("ALL HEADING 1 SECTIONS:")
    print("=" * 80)

    h1_count = 0
    chapter_numbers = []

    for i, para in enumerate(doc.paragraphs):
        if is_heading(para, 1):
            h1_count += 1
            text = para.text.strip()
            chapter_num = extract_chapter_number(text)

            if chapter_num:
                chapter_numbers.append(chapter_num)
                print(f"{h1_count:3d}. Para {i:4d}: Chapter {chapter_num:2d} - {text}")
            else:
                print(f"{h1_count:3d}. Para {i:4d}: (no number) - {text}")

    print("=" * 80)
    print(f"\nTotal Heading 1 sections: {h1_count}")
    print(f"Numbered chapters found: {len(chapter_numbers)}")

    if chapter_numbers:
        print(f"Chapter numbers: {sorted(chapter_numbers)}")
        print(f"Range: {min(chapter_numbers)} to {max(chapter_numbers)}")

        # Find missing numbers
        expected = set(range(1, 38))  # 1 to 37
        actual = set(chapter_numbers)
        missing = sorted(expected - actual)

        if missing:
            print(f"\n⚠️  Missing chapter numbers (expected 1-37): {missing}")
        else:
            print(f"\n✓ All chapters 1-37 present!")


if __name__ == '__main__':
    main()
