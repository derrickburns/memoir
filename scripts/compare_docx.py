#!/usr/bin/env python3
"""
Compare two DOCX files to check for differences.

This script compares the structure and content of two DOCX files
to determine if they are equivalent.
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import argparse


class DocxComparator:
    def __init__(self, file1, file2):
        """Initialize the comparator."""
        self.file1 = file1
        self.file2 = file2
        self.differences = []

    def has_image(self, run):
        """Check if a run contains an image."""
        drawing_elements = run._element.findall(f'.//{qn("a:blip")}',
                                                 namespaces=run._element.nsmap)
        return len(drawing_elements) > 0

    def count_images(self, doc):
        """Count total images in document."""
        count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if self.has_image(run):
                    count += 1
        return count

    def compare(self):
        """Compare the two documents."""
        print(f"Comparing:")
        print(f"  File 1: {self.file1}")
        print(f"  File 2: {self.file2}")
        print()

        doc1 = Document(self.file1)
        doc2 = Document(self.file2)

        # Compare paragraph counts
        para_count1 = len(doc1.paragraphs)
        para_count2 = len(doc2.paragraphs)

        print(f"Paragraph count:")
        print(f"  File 1: {para_count1}")
        print(f"  File 2: {para_count2}")

        if para_count1 != para_count2:
            diff = abs(para_count1 - para_count2)
            self.differences.append(f"Paragraph count differs by {diff}")
            print(f"  ❌ DIFFERENCE: {diff} paragraphs")
        else:
            print(f"  ✓ Match")
        print()

        # Compare image counts
        image_count1 = self.count_images(doc1)
        image_count2 = self.count_images(doc2)

        print(f"Image count:")
        print(f"  File 1: {image_count1}")
        print(f"  File 2: {image_count2}")

        if image_count1 != image_count2:
            diff = abs(image_count1 - image_count2)
            self.differences.append(f"Image count differs by {diff}")
            print(f"  ❌ DIFFERENCE: {diff} images")
        else:
            print(f"  ✓ Match")
        print()

        # Compare text content paragraph by paragraph
        print("Comparing paragraph content...")
        text_diffs = 0
        style_diffs = 0
        max_paras = min(para_count1, para_count2)

        for i in range(max_paras):
            para1 = doc1.paragraphs[i]
            para2 = doc2.paragraphs[i]

            # Compare text
            if para1.text != para2.text:
                text_diffs += 1
                if text_diffs <= 10:  # Show first 10 differences
                    print(f"  Para {i}: Text differs")
                    print(f"    File 1: {para1.text[:100]}...")
                    print(f"    File 2: {para2.text[:100]}...")

            # Compare styles
            style1 = para1.style.name if para1.style else None
            style2 = para2.style.name if para2.style else None

            if style1 != style2:
                style_diffs += 1
                if style_diffs <= 10:  # Show first 10 differences
                    print(f"  Para {i}: Style differs ({style1} vs {style2})")

        if text_diffs > 0:
            self.differences.append(f"{text_diffs} paragraphs have different text")
            print(f"\n  ❌ DIFFERENCE: {text_diffs} paragraphs have different text")
        else:
            print(f"  ✓ All paragraph text matches")

        if style_diffs > 0:
            self.differences.append(f"{style_diffs} paragraphs have different styles")
            print(f"  ❌ DIFFERENCE: {style_diffs} paragraphs have different styles")
        else:
            print(f"  ✓ All paragraph styles match")
        print()

        # Compare total text length
        text1 = '\n'.join([p.text for p in doc1.paragraphs])
        text2 = '\n'.join([p.text for p in doc2.paragraphs])

        print(f"Total text length:")
        print(f"  File 1: {len(text1)} characters")
        print(f"  File 2: {len(text2)} characters")

        if len(text1) != len(text2):
            diff = abs(len(text1) - len(text2))
            self.differences.append(f"Total text length differs by {diff} characters")
            print(f"  ❌ DIFFERENCE: {diff} characters")
        else:
            print(f"  ✓ Match")
        print()

        # Summary
        print("=" * 70)
        if len(self.differences) == 0:
            print("✓ FILES ARE IDENTICAL (idempotent operation)")
            return True
        else:
            print(f"❌ FILES DIFFER - {len(self.differences)} difference(s) found:")
            for i, diff in enumerate(self.differences, 1):
                print(f"  {i}. {diff}")
            return False


def main():
    parser = argparse.ArgumentParser(
        description="Compare two DOCX files for differences"
    )

    parser.add_argument('file1', help='First DOCX file')
    parser.add_argument('file2', help='Second DOCX file')

    args = parser.parse_args()

    # Check if files exist
    if not Path(args.file1).exists():
        print(f"ERROR: File not found: {args.file1}")
        sys.exit(1)

    if not Path(args.file2).exists():
        print(f"ERROR: File not found: {args.file2}")
        sys.exit(1)

    # Create comparator and run
    comparator = DocxComparator(args.file1, args.file2)
    result = comparator.compare()

    sys.exit(0 if result else 1)


if __name__ == '__main__':
    main()
