#!/usr/bin/env python3
"""
Renumber chapters sequentially to eliminate gaps.

This script renumbers chapters to be sequential (1, 2, 3, ...)
by updating both the file names and the chapter titles within the files.
"""

import os
import sys
import re
import shutil
from pathlib import Path
from docx import Document
import argparse


class ChapterRenumberer:
    def __init__(self, directory, dry_run=False):
        """Initialize the renumberer."""
        self.directory = Path(directory)
        self.dry_run = dry_run
        self.changes = []

    def extract_chapter_number(self, text):
        """Extract chapter number from text."""
        patterns = [
            r'[Cc]hapter\s+(\d+)',
            r'[Cc]h\.\s*(\d+)',
        ]

        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return int(match.group(1))

        return None

    def find_chapter_files(self):
        """Find all chapter files and their current numbers."""
        chapter_files = []

        for file_path in sorted(self.directory.glob("chapter_*.docx")):
            # Extract number from filename
            match = re.match(r'chapter_(\d+)', file_path.stem)
            if match:
                old_number = int(match.group(1))
                chapter_files.append((old_number, file_path))

        # Sort by old chapter number
        chapter_files.sort(key=lambda x: x[0])

        return chapter_files

    def analyze(self):
        """Analyze what changes would be made."""
        print(f"Analyzing chapters in: {self.directory}")
        print("=" * 80)

        chapter_files = self.find_chapter_files()

        if not chapter_files:
            print("No chapter files found!")
            return []

        # Create renumbering map
        renumbering_map = []
        new_number = 1

        for old_number, file_path in chapter_files:
            # Read the file to get the chapter title
            doc = Document(file_path)
            chapter_title = None

            # Find the first Heading 2 paragraph (should be chapter title)
            for para in doc.paragraphs:
                if para.style and para.style.name == 'Heading 2':
                    chapter_title = para.text.strip()
                    break

            if not chapter_title:
                chapter_title = file_path.stem

            renumbering_map.append({
                'old_number': old_number,
                'new_number': new_number,
                'file_path': file_path,
                'title': chapter_title,
            })

            new_number += 1

        # Display the mapping
        print(f"\nFound {len(renumbering_map)} chapters")
        print(f"Will renumber sequentially from 1 to {len(renumbering_map)}")
        print("\nRenumbering plan:")
        print("=" * 80)

        for item in renumbering_map:
            old_num = item['old_number']
            new_num = item['new_number']
            title = item['title']

            if old_num != new_num:
                print(f"Chapter {old_num:2d} → Chapter {new_num:2d}: {title[:60]}")
            else:
                print(f"Chapter {old_num:2d} (no change): {title[:60]}")

        return renumbering_map

    def renumber(self, renumbering_map):
        """Apply the renumbering."""
        if self.dry_run:
            print("\nDRY RUN - No changes will be made")
            return

        print("\nApplying renumbering...")
        print("=" * 80)

        # Process in reverse order to avoid conflicts
        for item in reversed(renumbering_map):
            old_number = item['old_number']
            new_number = item['new_number']
            old_path = item['file_path']

            if old_number == new_number:
                print(f"  Chapter {old_number}: No change needed")
                continue

            # Update the file content
            self._update_chapter_file(old_path, old_number, new_number)

            # Rename the file
            self._rename_chapter_file(old_path, old_number, new_number)

            print(f"  ✓ Renumbered Chapter {old_number} → Chapter {new_number}")

        print("\n✓ Renumbering complete!")

    def _update_chapter_file(self, file_path, old_number, new_number):
        """Update chapter number in the file content."""
        doc = Document(file_path)

        # Update all paragraphs that contain the old chapter number
        for para in doc.paragraphs:
            text = para.text

            # Replace "Chapter N" with "Chapter M" in the text
            new_text = re.sub(
                rf'\b[Cc]hapter\s+{old_number}\b',
                f'Chapter {new_number}',
                text
            )

            if new_text != text:
                # Update paragraph text while preserving formatting
                # We need to update each run
                if len(para.runs) > 0:
                    # Find which run contains the chapter number
                    for run in para.runs:
                        run_text = run.text
                        new_run_text = re.sub(
                            rf'\b[Cc]hapter\s+{old_number}\b',
                            f'Chapter {new_number}',
                            run_text
                        )
                        if new_run_text != run_text:
                            run.text = new_run_text

        # Save the updated document
        doc.save(file_path)

    def _rename_chapter_file(self, old_path, old_number, new_number):
        """Rename chapter file and associated image directory."""
        # Rename the DOCX file
        old_stem = old_path.stem
        new_stem = re.sub(rf'^chapter_{old_number:02d}', f'chapter_{new_number:02d}', old_stem)
        new_path = old_path.parent / f"{new_stem}.docx"

        # Use a temporary name to avoid conflicts
        temp_path = old_path.parent / f"temp_renumber_{new_number:02d}.docx"
        shutil.move(str(old_path), str(temp_path))
        shutil.move(str(temp_path), str(new_path))

        # Rename the image directory if it exists
        old_img_dir = old_path.parent / f"chapter_{old_number:02d}_images"
        new_img_dir = old_path.parent / f"chapter_{new_number:02d}_images"

        if old_img_dir.exists():
            temp_img_dir = old_path.parent / f"temp_renumber_{new_number:02d}_images"
            shutil.move(str(old_img_dir), str(temp_img_dir))
            shutil.move(str(temp_img_dir), str(new_img_dir))


def main():
    parser = argparse.ArgumentParser(
        description="Renumber chapters sequentially to eliminate gaps",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Analyze what would change
  %(prog)s output --dry-run

  # Renumber chapters
  %(prog)s output

This will:
  1. Find all chapter files (chapter_*.docx)
  2. Renumber them sequentially (1, 2, 3, ...)
  3. Update chapter titles in the files
  4. Rename files and image directories
        """
    )

    parser.add_argument('directory', help='Directory containing chapter files')
    parser.add_argument('-d', '--dry-run', action='store_true',
                        help='Show what would change without making changes')

    args = parser.parse_args()

    # Check directory exists
    if not os.path.exists(args.directory):
        print(f"ERROR: Directory not found: {args.directory}")
        sys.exit(1)

    # Create renumberer and run
    renumberer = ChapterRenumberer(args.directory, dry_run=args.dry_run)
    renumbering_map = renumberer.analyze()

    if renumbering_map:
        print("\n" + "=" * 80)
        renumberer.renumber(renumbering_map)
    else:
        print("\nNo chapters to renumber.")


if __name__ == '__main__':
    main()
